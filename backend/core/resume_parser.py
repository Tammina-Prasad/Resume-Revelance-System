"""
Resume Parser Module
Handles extraction and cleaning of text from PDF and DOCX resume files.
"""

import os
import re
import logging
from typing import Optional, Dict, Any
from pathlib import Path

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """
    A class to handle parsing of resume files in PDF and DOCX formats.
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Main method to parse a resume file and extract text.
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            Dict[str, Any]: Dictionary containing extracted text and metadata
        """
        try:
            # Validate file path
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file extension
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text based on file type
            if file_extension == '.pdf':
                raw_text = self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                raw_text = self._extract_docx_text(file_path)
            
            # Clean and standardize text
            cleaned_text = self._clean_text(raw_text)
            
            # Extract basic information
            parsed_data = {
                'raw_text': raw_text,
                'cleaned_text': cleaned_text,
                'file_path': file_path,
                'file_type': file_extension,
                'word_count': len(cleaned_text.split()),
                'char_count': len(cleaned_text),
                'sections': self._identify_sections(cleaned_text)
            }
            
            logger.info(f"Successfully parsed resume: {file_path}")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from PDF file using multiple methods for better accuracy.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text
        """
        text = ""
        
        try:
            # Method 1: Using pdfplumber (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber fails, try PyPDF2
            if not text.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            # Fallback to PyPDF2 if pdfplumber fails
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {str(e2)}")
                raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _clean_text(self, raw_text: str) -> str:
        """
        Clean and standardize extracted text.
        
        Args:
            raw_text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not raw_text:
            return ""
        
        # Remove extra whitespace and newlines
        text = re.sub(r'\s+', ' ', raw_text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s@.,()-]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        # Convert to lowercase for consistency (optional - might want to preserve case)
        # text = text.lower()
        
        return text
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify common resume sections using pattern matching.
        
        Args:
            text (str): Cleaned text
            
        Returns:
            Dict[str, str]: Dictionary of identified sections
        """
        sections = {}
        
        # Common section headers
        section_patterns = {
            'contact': r'(contact|email|phone|address)',
            'summary': r'(summary|profile|objective|about)',
            'experience': r'(experience|employment|work history|professional)',
            'education': r'(education|academic|qualification|degree)',
            'skills': r'(skills|technical|competencies|expertise)',
            'projects': r'(projects|portfolio|work samples)',
            'certifications': r'(certification|certificate|license)',
            'achievements': r'(achievement|award|recognition|honor)'
        }
        
        # Split text into potential sections
        lines = text.split('\n')
        current_section = 'general'
        sections[current_section] = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line matches any section header
            section_found = False
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line.lower()) and len(line.split()) <= 5:
                    current_section = section_name
                    sections[current_section] = ""
                    section_found = True
                    break
            
            if not section_found:
                sections[current_section] += line + " "
        
        # Clean up sections
        for section in sections:
            sections[section] = sections[section].strip()
        
        return sections


def extract_contact_info(text: str) -> Dict[str, str]:
    """
    Extract contact information from resume text.
    
    Args:
        text (str): Resume text
        
    Returns:
        Dict[str, str]: Extracted contact information
    """
    contact_info = {}
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        contact_info['email'] = emails[0]
    
    # Phone pattern (multiple formats)
    phone_pattern = r'(\+?\d{1,4}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    if phones:
        contact_info['phone'] = ''.join(phones[0])
    
    # LinkedIn pattern
    linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
    linkedin = re.findall(linkedin_pattern, text.lower())
    if linkedin:
        contact_info['linkedin'] = linkedin[0]
    
    return contact_info


def extract_education_info(text: str) -> list:
    """
    Extract education information from resume text.
    
    Args:
        text (str): Resume text
        
    Returns:
        list: List of education entries
    """
    education_keywords = [
        'bachelor', 'master', 'phd', 'doctorate', 'degree', 'university', 
        'college', 'institute', 'b.tech', 'm.tech', 'mba', 'ms', 'bs'
    ]
    
    education_info = []
    lines = text.lower().split('\n')
    
    for line in lines:
        if any(keyword in line for keyword in education_keywords):
            education_info.append(line.strip())
    
    return education_info


# Example usage and testing
if __name__ == "__main__":
    parser = ResumeParser()
    
    # Test with a sample file (uncomment when testing)
    # result = parser.parse_resume("path/to/sample_resume.pdf")
    # print(f"Extracted text length: {len(result['cleaned_text'])}")
    # print(f"Sections found: {list(result['sections'].keys())}")